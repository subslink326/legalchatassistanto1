"""
Generate a Markdown outline (and optional .docx) for the requested motion/brief.
"""

from __future__ import annotations
from pathlib import Path
import datetime
import markdown
from uuid import UUID
from typing import List

from pydantic import BaseModel, Field
from jinja2 import Environment, FileSystemLoader, select_autoescape
import docx

TEMPLATE_DIR = Path(__file__).parent / "templates"
env = Environment(
    loader=FileSystemLoader(TEMPLATE_DIR),
    autoescape=select_autoescape(),
    trim_blocks=True,
    lstrip_blocks=True,
)

# ---------- Pydantic ----------
class BriefRequest(BaseModel):
    case_title: str = Field(..., example="United States v. Doe, No. 21‑cr‑42")
    motion_type: str = Field(..., example="Fed. R. Crim. P. 33 Motion")
    issues: List[str]
    party: str = Field("Defendant", example="Defendant")
    prayer: str = Field("a new trial", example="a new trial")
    generate_docx: bool = False


class BriefResponse(BaseModel):
    outline_markdown: str
    docx_filename: str | None = None


# ---------- Helper ----------
def _render_md(req: BriefRequest) -> str:
    tpl_name = "fed_rule33.md"  # later map motion_type -> template file
    tpl = env.get_template(tpl_name)
    return tpl.render(
        case_title=req.case_title,
        motion_type=req.motion_type,
        issues=req.issues,
        party=req.party,
        prayer=req.prayer,
        date=datetime.date.today().isoformat(),
    )


def _md_to_docx(md_text: str, out_path: Path) -> None:
    """Very naive Markdown → docx via python‑docx; keeps only headings / paragraphs."""
    html = markdown.markdown(md_text)
    doc = docx.Document()
    for block in html.split("</p>"):
        clean = (
            block.replace("<p>", "")
            .replace("<strong>", "")
            .replace("</strong>", "")
            .strip()
        )
        if not clean:
            continue
        if clean.startswith("<h"):
            doc.add_heading(clean.split(">")[1], level=2)
        else:
            doc.add_paragraph(clean)
    doc.save(out_path)


# ---------- Main entry ----------
async def generate_brief(req: BriefRequest) -> BriefResponse:
    md_text = _render_md(req)
    docx_file = None
    if req.generate_docx:
        out_path = Path("storage") / f"draft_{datetime.datetime.now().timestamp()}.docx"
        _md_to_docx(md_text, out_path)
        docx_file = str(out_path)
    return BriefResponse(outline_markdown=md_text, docx_filename=docx_file)
